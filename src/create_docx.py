import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MAX_THERMAL_IMAGES_IN_REPORT = 8

ACCENT = "1F4E79"
ACCENT_LIGHT = "D9EAF7"
GOLD = "F2B705"
LIGHT_GRAY = "F2F2F2"
DARK_TEXT = "222222"
WARNING = "FFF2CC"
HIGH = "F4CCCC"
MODERATE = "D9EAD3"


# ---------------------------------------------------------
# Basic styling helpers
# ---------------------------------------------------------

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_borders(table, color="D9D9D9"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color)


def set_paragraph_font(paragraph, size=10, bold=False, color=DARK_TEXT):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_spacer(doc, height=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(height)


def add_section_title(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)

    add_spacer(doc, 4)


def add_label_value_table(doc, data):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for label, value in data:
        row = table.add_row()
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        set_cell_shading(label_cell, LIGHT_GRAY)

        label_cell.text = str(label)
        value_cell.text = str(value)

        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for p in cell.paragraphs:
                set_paragraph_font(p, size=9)

        for p in label_cell.paragraphs:
            set_paragraph_font(p, size=9, bold=True)

    add_spacer(doc, 8)
    return table


def get_severity_color(severity):
    severity_lower = str(severity).lower()

    if "high" in severity_lower:
        return HIGH

    if "moderate" in severity_lower:
        return MODERATE

    return WARNING


# ---------------------------------------------------------
# Header / Footer / Document setup
# ---------------------------------------------------------

def setup_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles

    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK_TEXT)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.font.bold = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def add_header_footer(doc):
    section = doc.sections[0]

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Detailed Diagnostic Report"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Generated DDR | AI-assisted evidence-based report"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")


# ---------------------------------------------------------
# Cover page
# ---------------------------------------------------------

def add_cover_page(doc, evidence):
    property_details = evidence.get("property_details", {})

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)

    run = title.add_run("DETAILED DIAGNOSTIC REPORT")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subtitle.add_run("Property Inspection & Thermal Evidence Summary")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("555555")

    add_spacer(doc, 20)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_LIGHT)
    set_cell_border(cell, color="B7D7EA")

    lines = [
        ("Property", property_details.get("flat_no", "Not Available")),
        ("Property Type", property_details.get("property_type", "Not Available")),
        ("Inspection Date", property_details.get("inspection_date_time", "Not Available")),
        ("Inspected By", property_details.get("inspected_by", "Not Available")),
    ]

    for label, value in lines:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)

        r2 = p.add_run(str(value))
        r2.font.size = Pt(11)

    add_spacer(doc, 28)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Inches(0.7)
    note.paragraph_format.right_indent = Inches(0.7)

    run = note.add_run(
        "This report is generated from the provided inspection report and thermal report. "
        "Missing or unclear details are marked as Not Available."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_page_break()


# ---------------------------------------------------------
# Main report sections
# ---------------------------------------------------------

def add_executive_summary(doc, evidence):
    property_details = evidence.get("property_details", {})
    observations = evidence.get("observations", [])
    thermal_readings = evidence.get("thermal_readings", [])

    add_section_title(doc, "1. Property Issue Summary")

    summary_text = (
        f"The inspection of {property_details.get('flat_no', 'Not Available')} "
        f"was conducted on {property_details.get('inspection_date_time', 'Not Available')} "
        f"by {property_details.get('inspected_by', 'Not Available')}. "
        f"The inspection identified {len(observations)} key affected areas involving dampness, "
        f"leakage, moisture movement, tile joint gaps, possible plumbing issues, and external wall cracks. "
        f"Thermal evidence is available as reference evidence with {len(thermal_readings)} thermal references. "
        "Exact room-wise mapping of every thermal image is Not Available."
    )

    p = doc.add_paragraph(summary_text)
    p.paragraph_format.space_after = Pt(8)

    property_table_data = [
        ("Flat / Unit", property_details.get("flat_no", "Not Available")),
        ("Property Type", property_details.get("property_type", "Not Available")),
        ("Floors", property_details.get("floors", "Not Available")),
        ("Inspection Date & Time", property_details.get("inspection_date_time", "Not Available")),
        ("Inspected By", property_details.get("inspected_by", "Not Available")),
        ("Previous Structural Audit", property_details.get("previous_structural_audit_done", "Not Available")),
        ("Previous Repair Work", property_details.get("previous_repair_work_done", "Not Available")),
    ]

    add_label_value_table(doc, property_table_data)


def add_area_observations_table(doc, evidence):
    observations = evidence.get("observations", [])

    add_section_title(doc, "2. Area-wise Observations")

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    headers = ["No.", "Area", "Observation", "Probable Source", "Severity"]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, ACCENT)

        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    for obs in observations:
        row = table.add_row().cells

        values = [
            obs.get("id", "Not Available"),
            obs.get("impacted_area", "Not Available"),
            obs.get("negative_side_observation", "Not Available"),
            obs.get("probable_positive_side_source", "Not Available"),
            obs.get("severity", "Not Available"),
        ]

        for idx, value in enumerate(values):
            row[idx].text = str(value)

            for p in row[idx].paragraphs:
                set_paragraph_font(p, size=8)

        set_cell_shading(row[4], get_severity_color(obs.get("severity", "")))

    add_spacer(doc, 10)


def add_root_cause_and_severity(doc, evidence):
    add_section_title(doc, "3. Probable Root Cause")

    causes = [
        "Gaps between tile joints in Common Bathroom and Master Bedroom Bathroom may be allowing water to penetrate below tiled surfaces.",
        "Concealed plumbing leakage is possible based on checklist findings.",
        "Gaps around Nahani trap joints and plumbing outlet joints may be contributing to moisture movement.",
        "External wall cracks near the Master Bedroom may be allowing water ingress.",
        "Exact moisture travel path is Not Available for every affected area."
    ]

    for cause in causes:
        doc.add_paragraph(cause, style="List Bullet")

    add_section_title(doc, "4. Severity Assessment with Reasoning")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = ["Area", "Severity", "Reasoning"]

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, ACCENT)

        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    for obs in evidence.get("observations", []):
        row = table.add_row().cells
        row[0].text = obs.get("impacted_area", "Not Available")
        row[1].text = obs.get("severity", "Not Available")
        row[2].text = obs.get("reasoning", "Not Available")

        set_cell_shading(row[1], get_severity_color(obs.get("severity", "")))

        for cell in row:
            for p in cell.paragraphs:
                set_paragraph_font(p, size=8)

    add_spacer(doc, 10)


def add_actions_and_missing_info(doc, evidence):
    add_section_title(doc, "5. Recommended Actions")

    actions = evidence.get("recommended_actions", [])

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    table.rows[0].cells[0].text = "Priority"
    table.rows[0].cells[1].text = "Recommended Action"

    for cell in table.rows[0].cells:
        set_cell_shading(cell, ACCENT)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    for idx, action in enumerate(actions, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = action

        for cell in row:
            for p in cell.paragraphs:
                set_paragraph_font(p, size=9)

    add_section_title(doc, "6. Additional Notes")

    notes = [
        "Visual images are included under the Area-wise Visual Evidence section.",
        "Thermal images are included as reference evidence.",
        "Exact room-wise mapping of every thermal image is Not Available because the source thermal document does not provide explicit area labels.",
        "The report avoids unsupported conclusions and uses probable root cause language."
    ]

    for note in notes:
        doc.add_paragraph(note, style="List Bullet")

    add_section_title(doc, "7. Missing or Unclear Information")

    missing_items = evidence.get("missing_or_unclear_information", [])

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Status"

    for cell in table.rows[0].cells:
        set_cell_shading(cell, ACCENT)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    for item in missing_items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = "Not Available"
        set_cell_shading(row[1], WARNING)

        for cell in row:
            for p in cell.paragraphs:
                set_paragraph_font(p, size=9)


# ---------------------------------------------------------
# Images
# ---------------------------------------------------------

def should_skip_image_record(image: dict) -> bool:
    image_type = image.get("type", "")

    if image_type == "thermal_page":
        return False

    width = int(image.get("width", 0))
    height = int(image.get("height", 0))

    if width <= 0 or height <= 0:
        return False

    if width < 120 or height < 120:
        return True

    aspect_ratio = width / height

    if aspect_ratio > 4.0:
        return True

    if aspect_ratio < 0.20:
        return True

    return False


def add_image_to_cell(cell, image_path, width_inches=2.35):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_image_grid(doc, images, cols=2, image_width=2.35):
    valid_images = []

    for image in images:
        if should_skip_image_record(image):
            continue

        image_path = image.get("image_path")

        if not image_path:
            continue

        image_file = Path(image_path)

        if image_file.exists():
            valid_images.append(image)

    if not valid_images:
        doc.add_paragraph("Image Not Available")
        return

    for i in range(0, len(valid_images), cols):
        row_images = valid_images[i:i + cols]

        table = doc.add_table(rows=2, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for col_idx in range(cols):
            img_cell = table.cell(0, col_idx)
            cap_cell = table.cell(1, col_idx)

            if col_idx < len(row_images):
                image = row_images[col_idx]
                image_file = Path(image["image_path"])

                try:
                    add_image_to_cell(img_cell, image_file, width_inches=image_width)
                    cap_cell.text = f"Source Page {image.get('page_number', 'Not Available')}"
                    for p in cap_cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_paragraph_font(p, size=8, color="666666")
                except Exception:
                    img_cell.text = "Image could not be inserted"
            else:
                img_cell.text = ""
                cap_cell.text = ""

        add_spacer(doc, 6)


def add_area_visual_evidence(doc, evidence):
    doc.add_page_break()
    add_section_title(doc, "Area-wise Visual Evidence")

    for obs in evidence.get("observations", []):
        area = obs.get("impacted_area", "Not Available")
        obs_id = obs.get("id", "")
        observation = obs.get("negative_side_observation", "Not Available")
        source = obs.get("probable_positive_side_source", "Not Available")
        severity = obs.get("severity", "Not Available")

        doc.add_heading(f"{obs_id}. {area}", level=2)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(info_table)

        rows = [
            ("Observation", observation),
            ("Probable Source", source),
            ("Severity", severity)
        ]

        for idx, (label, value) in enumerate(rows):
            info_table.rows[idx].cells[0].text = label
            info_table.rows[idx].cells[1].text = value

            set_cell_shading(info_table.rows[idx].cells[0], LIGHT_GRAY)

            if label == "Severity":
                set_cell_shading(info_table.rows[idx].cells[1], get_severity_color(severity))

            for cell in info_table.rows[idx].cells:
                for p in cell.paragraphs:
                    set_paragraph_font(p, size=9)

            for p in info_table.rows[idx].cells[0].paragraphs:
                set_paragraph_font(p, size=9, bold=True)

        add_spacer(doc, 4)

        doc.add_heading("Related Inspection Images", level=3)
        add_image_grid(
            doc,
            obs.get("related_images", []),
            cols=2,
            image_width=2.35
        )


# ---------------------------------------------------------
# Thermal evidence
# ---------------------------------------------------------

def parse_temp(value):
    try:
        return float(str(value).replace("°C", "").strip())
    except Exception:
        return None


def add_thermal_summary(doc, thermal_readings):
    hotspots = [
        parse_temp(reading.get("hotspot"))
        for reading in thermal_readings
    ]

    coldspots = [
        parse_temp(reading.get("coldspot"))
        for reading in thermal_readings
    ]

    hotspots = [x for x in hotspots if x is not None]
    coldspots = [x for x in coldspots if x is not None]

    table_data = [
        ("Total Thermal References", len(thermal_readings)),
        ("Hotspot Range", f"{min(hotspots)} °C to {max(hotspots)} °C" if hotspots else "Not Available"),
        ("Coldspot Range", f"{min(coldspots)} °C to {max(coldspots)} °C" if coldspots else "Not Available"),
        ("Exact Room-wise Mapping", "Not Available")
    ]

    add_label_value_table(doc, table_data)


def add_thermal_evidence(doc, evidence):
    doc.add_page_break()
    add_section_title(doc, "Thermal Evidence")

    thermal_readings = evidence.get("thermal_readings", [])

    if not thermal_readings:
        doc.add_paragraph("Thermal evidence: Not Available")
        return

    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run(
        "Thermal images are included as reference evidence. Exact mapping between every thermal image "
        "and impacted area is Not Available unless explicitly mentioned in the source documents."
    )

    add_thermal_summary(doc, thermal_readings)

    doc.add_heading("Selected Thermal References", level=2)

    selected_readings = thermal_readings[:MAX_THERMAL_IMAGES_IN_REPORT]

    p = doc.add_paragraph(
        f"The first {len(selected_readings)} thermal references are included below. "
        f"Total thermal references available: {len(thermal_readings)}."
    )
    set_paragraph_font(p, size=9)

    for reading in selected_readings:
        doc.add_heading(f"Thermal Reference - Page {reading.get('page_number', 'Not Available')}", level=3)

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        details = [
            ("Thermal Image", reading.get("thermal_image", "Not Available")),
            ("Hotspot", reading.get("hotspot", "Not Available")),
            ("Coldspot", reading.get("coldspot", "Not Available")),
            ("Emissivity", reading.get("emissivity", "Not Available")),
            ("Reflected Temperature", reading.get("reflected_temperature", "Not Available")),
            ("Device", reading.get("device", "Not Available")),
            ("Area Mapping", reading.get("area_mapping", "Not Available")),
        ]

        for label, value in details:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

            set_cell_shading(row[0], LIGHT_GRAY)

            if label == "Area Mapping":
                set_cell_shading(row[1], WARNING)

            for cell in row:
                for para in cell.paragraphs:
                    set_paragraph_font(para, size=8)

            for para in row[0].paragraphs:
                set_paragraph_font(para, size=8, bold=True)

        add_spacer(doc, 4)

        add_image_grid(
            doc,
            reading.get("images", []),
            cols=1,
            image_width=4.7
        )


# ---------------------------------------------------------
# Main
# ---------------------------------------------------------

def create_docx():
    doc = Document()

    evidence_path = Path("output/evidence.json")

    if not evidence_path.exists():
        raise FileNotFoundError("output/evidence.json not found. Run build_evidence.py first.")

    with open(evidence_path, "r", encoding="utf-8") as f:
        evidence = json.load(f)

    setup_document_styles(doc)
    add_header_footer(doc)

    add_cover_page(doc, evidence)
    add_executive_summary(doc, evidence)
    add_area_observations_table(doc, evidence)
    add_root_cause_and_severity(doc, evidence)
    add_actions_and_missing_info(doc, evidence)
    add_area_visual_evidence(doc, evidence)
    add_thermal_evidence(doc, evidence)

    output_path = Path("output/generated_ddr.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(output_path)

    print(f"Professionally formatted DOCX created: {output_path}")


if __name__ == "__main__":
    create_docx()